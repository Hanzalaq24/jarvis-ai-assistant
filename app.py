from flask import Flask, jsonify, request, render_template
from flask_cors import CORS
import webbrowser
import subprocess
import os
from datetime import datetime
from threading import Thread, Lock

# Text-to-Speech
try:
    import pyttsx3
    TTS_AVAILABLE = True
except Exception as e:
    TTS_AVAILABLE = False
    print(f"Text-to-speech not available: {e}")

# AI Assistant
try:
    from ai_assistant import get_ai_response
    AI_AVAILABLE = True
    print("🤖 AI Assistant loaded successfully!")
except ImportError as e:
    AI_AVAILABLE = False
    print(f"⚠️ AI Assistant not available: {e}")

# System utilities
import psutil
import pyautogui
from pathlib import Path

# Photo capture
try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    print("OpenCV not available - photo capture disabled")

# Volume control
try:
    from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
    from comtypes import CLSCTX_ALL
    PYCAW_AVAILABLE = True
except Exception as e:
    PYCAW_AVAILABLE = False
    print(f"Pycaw not available - volume control disabled: {e}")

# Improved File Operations
try:
    from improved_file_operations import file_ops
    from improved_command_processor import command_processor
    IMPROVED_FILE_OPS_AVAILABLE = True
    print("✅ Improved file operations loaded successfully!")
except ImportError as e:
    IMPROVED_FILE_OPS_AVAILABLE = False
    print(f"⚠️ Improved file operations not available: {e}")

# Multi-language support imports
try:
    from googletrans import Translator
    from langdetect import detect
    TRANSLATION_AVAILABLE = True
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("Translation libraries not available - multi-language support disabled")

# Audio Processing for Song Recognition
try:
    import pyaudio
    import librosa
    import numpy as np
    import soundfile as sf
    from scipy import signal
    AUDIO_PROCESSING_AVAILABLE = True
    print("🎵 Audio processing libraries loaded successfully!")
except ImportError as e:
    AUDIO_PROCESSING_AVAILABLE = False
    print(f"🎵 Audio processing not available: {e}")
    print("Install with: pip install pyaudio librosa numpy scipy soundfile")

# File operations imports
try:
    import send2trash
    import shutil
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    FILE_OPS_AVAILABLE = True
except ImportError:
    FILE_OPS_AVAILABLE = False
    print("File operations libraries not available - some features disabled")

# Advanced features imports
try:
    import requests
    from bs4 import BeautifulSoup
    import urllib.parse
    ADVANCED_SEARCH_AVAILABLE = True
except ImportError:
    ADVANCED_SEARCH_AVAILABLE = False
    print("Advanced search features not available")

try:
    from plyer import notification
    NOTIFICATIONS_AVAILABLE = True
except ImportError:
    NOTIFICATIONS_AVAILABLE = False
    print("Notifications not available")

import glob
import mimetypes
import json
import platform
import sys
import re

BASE_DIR = Path(__file__).parent.resolve()

# Detect operating system
CURRENT_OS = platform.system().lower()
IS_WINDOWS = CURRENT_OS == 'windows'
IS_MACOS = CURRENT_OS == 'darwin'
IS_LINUX = CURRENT_OS == 'linux'

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)
app.secret_key = 'your-secret-key-123'

# Speech Engine with better initialization
engine = None
if TTS_AVAILABLE:
    try:
        engine = pyttsx3.init()
        
        # Test if engine is working
        voices = engine.getProperty('voices')
        if not voices:
            raise Exception("No voices available")
        
        # Set initial properties
        engine.setProperty('rate', 150)
        engine.setProperty('volume', 0.9)
        
        # Find and set the best voice
        best_voice = None
        for voice in voices:
            voice_name = voice.name.lower() if hasattr(voice, 'name') else ''
            if any(keyword in voice_name for keyword in ['zira', 'hazel', 'samantha']):
                best_voice = voice.id
                break
            elif 'english' in voice_name and not best_voice:
                best_voice = voice.id
        
        if best_voice:
            engine.setProperty('voice', best_voice)
        elif len(voices) > 1:
            engine.setProperty('voice', voices[1].id)
        
        # Test the engine
        engine.say("TTS initialized")
        engine.runAndWait()
        
        print(f"✅ TTS Engine initialized successfully with {len(voices)} voices")
        
    except Exception as e:
        TTS_AVAILABLE = False
        engine = None
        print(f"❌ Failed to initialize TTS engine: {e}")
        print("TTS will be disabled - responses will be text-only")

# Multi-language setup
if TRANSLATION_AVAILABLE:
    translator = Translator()
else:
    translator = None

# Language mappings
LANGUAGE_CODES = {
    'english': 'en',
    'hindi': 'hi', 
    'gujarati': 'gu'
}

# Multi-language responses
RESPONSES = {
    'en': {
        'greeting': 'Hello! I am JARVIS, your AI assistant.',
        'time_prefix': 'The current time is',
        'date_prefix': "Today's date is",
        'battery_prefix': 'Battery at',
        'opening': 'Opening',
        'not_understood': "I didn't understand that command",
        'reminder_set': 'Reminder set for',
        'photo_captured': 'Photo captured and saved',
        'volume_muted': 'Volume muted',
        'volume_unmuted': 'Volume unmuted',
        'volume_increased': 'Volume increased',
        'volume_decreased': 'Volume decreased'
    },
    'hi': {
        'greeting': 'नमस्ते! मैं जार्विस हूं, आपका AI सहायक।',
        'time_prefix': 'वर्तमान समय है',
        'date_prefix': 'आज की तारीख है',
        'battery_prefix': 'बैटरी',
        'opening': 'खोल रहा हूं',
        'not_understood': 'मुझे यह कमांड समझ नहीं आया',
        'reminder_set': 'रिमाइंडर सेट किया गया',
        'photo_captured': 'फोटो कैप्चर और सेव किया गया',
        'volume_muted': 'वॉल्यूम म्यूट किया गया',
        'volume_unmuted': 'वॉल्यूम अनम्यूट किया गया',
        'volume_increased': 'वॉल्यूम बढ़ाया गया',
        'volume_decreased': 'वॉल्यूम कम किया गया'
    },
    'gu': {
        'greeting': 'નમસ્તે! હું જાર્વિસ છું, તમારો AI સહાયક.',
        'time_prefix': 'વર્તમાન સમય છે',
        'date_prefix': 'આજની તારીખ છે',
        'battery_prefix': 'બેટરી',
        'opening': 'ખોલી રહ્યો છું',
        'not_understood': 'મને આ કમાન્ડ સમજાયો નહીં',
        'reminder_set': 'રિમાઇન્ડર સેટ કર્યું',
        'photo_captured': 'ફોટો કેપ્ચર અને સેવ કર્યો',
        'volume_muted': 'વોલ્યુમ મ્યુટ કર્યું',
        'volume_unmuted': 'વોલ્યુમ અનમ્યુટ કર્યું',
        'volume_increased': 'વોલ્યુમ વધાર્યું',
        'volume_decreased': 'વોલ્યુમ ઘટાડ્યું'
    }
}

WEBSITE_MAP = {
    'google': 'https://www.google.com',
    'youtube': 'https://www.youtube.com',
    'facebook': 'https://www.facebook.com',
    'twitter': 'https://www.twitter.com',
    'instagram': 'https://www.instagram.com',
    'linkedin': 'https://www.linkedin.com',
    'github': 'https://www.github.com',
    'amazon': 'https://www.amazon.com',
    'netflix': 'https://www.netflix.com',
    'wikipedia': 'https://www.wikipedia.org',
    'gmail': 'https://mail.google.com',
    'outlook': 'https://outlook.live.com',
    'reddit': 'https://www.reddit.com',
    'stackoverflow': 'https://stackoverflow.com',
    'spotify': 'https://open.spotify.com',
    'discord': 'https://discord.com',
    'whatsapp': 'https://web.whatsapp.com',
    'telegram': 'https://web.telegram.org',
}

# Enhanced search URL patterns
SEARCH_PATTERNS = {
    'google': 'https://www.google.com/search?q={}',
    'youtube': 'https://www.youtube.com/results?search_query={}',
    'wikipedia': 'https://en.wikipedia.org/wiki/{}',
    'amazon': 'https://www.amazon.com/s?k={}',
    'reddit': 'https://www.reddit.com/search/?q={}',
    'stackoverflow': 'https://stackoverflow.com/search?q={}',
    'github': 'https://github.com/search?q={}',
    'spotify': 'https://open.spotify.com/search/{}',
    'netflix': 'https://www.netflix.com/search?q={}',
}

# Cross-platform application mapping with fallbacks
if IS_WINDOWS:
    APP_MAP = {
        'chrome': ['C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe', 
                  'C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe'],
        'edge': ['C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe',
                'C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe'],
        'firefox': ['C:\\Program Files\\Mozilla Firefox\\firefox.exe',
                   'C:\\Program Files (x86)\\Mozilla Firefox\\firefox.exe'],
        'notepad': ['notepad.exe'],
        'calculator': ['calc.exe'],
        'word': ['winword.exe', 'C:\\Program Files\\Microsoft Office\\root\\Office16\\WINWORD.EXE'],
        'excel': ['excel.exe', 'C:\\Program Files\\Microsoft Office\\root\\Office16\\EXCEL.EXE'],
        'powerpoint': ['powerpnt.exe', 'C:\\Program Files\\Microsoft Office\\root\\Office16\\POWERPNT.EXE'],
        'spotify': ['spotify.exe'],
        'vscode': ['code.exe', 'C:\\Users\\%USERNAME%\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe'],
        'photoshop': ['photoshop.exe'],
        'steam': ['steam.exe'],
        'discord': ['discord.exe'],
        'telegram': ['telegram.exe'],
        'whatsapp': ['whatsapp.exe'],
        'cmd': ['cmd.exe'],
        'powershell': ['powershell.exe'],
        'paint': ['mspaint.exe'],
        'explorer': ['explorer.exe'],
    }
elif IS_MACOS:
    APP_MAP = {
        'chrome': ['/Applications/Google Chrome.app'],
        'safari': ['/Applications/Safari.app', '/System/Applications/Safari.app'],
        'firefox': ['/Applications/Firefox.app'],
        'textedit': ['/Applications/TextEdit.app', '/System/Applications/TextEdit.app'],
        'calculator': ['/Applications/Calculator.app', '/System/Applications/Calculator.app'],
        'word': ['/Applications/Microsoft Word.app'],
        'excel': ['/Applications/Microsoft Excel.app'],
        'powerpoint': ['/Applications/Microsoft PowerPoint.app'],
        'spotify': ['/Applications/Spotify.app'],
        'vscode': ['/Applications/Visual Studio Code.app'],
        'photoshop': ['/Applications/Adobe Photoshop 2024/Adobe Photoshop 2024.app',
                     '/Applications/Adobe Photoshop 2023/Adobe Photoshop 2023.app'],
        'steam': ['/Applications/Steam.app'],
        'discord': ['/Applications/Discord.app'],
        'telegram': ['/Applications/Telegram.app'],
        'whatsapp': ['/Applications/WhatsApp.app'],
        'finder': ['/System/Library/CoreServices/Finder.app'],
        'terminal': ['/Applications/Utilities/Terminal.app', '/System/Applications/Utilities/Terminal.app'],
        'preview': ['/Applications/Preview.app', '/System/Applications/Preview.app'],
        'mail': ['/Applications/Mail.app', '/System/Applications/Mail.app'],
        'notes': ['/Applications/Notes.app', '/System/Applications/Notes.app'],
    }
else:  # Linux
    APP_MAP = {
        'chrome': ['google-chrome', 'google-chrome-stable', 'chromium-browser'],
        'firefox': ['firefox'],
        'gedit': ['gedit', 'nano', 'vim'],
        'calculator': ['gnome-calculator', 'kcalc', 'galculator'],
        'spotify': ['spotify'],
        'vscode': ['code', 'codium'],
        'steam': ['steam'],
        'discord': ['discord'],
        'telegram': ['telegram-desktop', 'telegram'],
        'nautilus': ['nautilus', 'dolphin', 'thunar'],
        'terminal': ['gnome-terminal', 'konsole', 'xterm'],
    }

reminders = []

def detect_language(text):
    """Detect the language of input text with improved accuracy"""
    if not TRANSLATION_AVAILABLE:
        return 'en'
    
    try:
        # Check for specific Hindi/Gujarati keywords first
        hindi_keywords = ['समय', 'बैटरी', 'मदद', 'खोलो', 'बंद', 'चालू', 'फाइल', 'फोल्डर', 'क्या', 'है', 'करो', 'बनाओ']
        gujarati_keywords = ['સમય', 'બેટરી', 'મદદ', 'ખોલો', 'બંધ', 'ચાલુ', 'ફાઇલ', 'ફોલ્ડર', 'શું', 'છે', 'કરો', 'બનાવો']
        
        text_lower = text.lower()
        
        # Check for Hindi keywords
        for keyword in hindi_keywords:
            if keyword in text:
                return 'hi'
        
        # Check for Gujarati keywords
        for keyword in gujarati_keywords:
            if keyword in text:
                return 'gu'
        
        # Use automatic detection as fallback
        detected = detect(text)
        if detected in ['hi', 'gu', 'en']:
            return detected
        return 'en'  # Default to English
    except:
        return 'en'

def translate_text(text, target_lang):
    """Translate text to target language"""
    if not TRANSLATION_AVAILABLE or target_lang == 'en':
        return text
    
    try:
        result = translator.translate(text, dest=target_lang)
        return result.text
    except:
        return text

def get_response_text(key, lang_code, *args):
    """Get localized response text"""
    if lang_code not in RESPONSES:
        lang_code = 'en'
    
    if key in RESPONSES[lang_code]:
        response = RESPONSES[lang_code][key]
        if args:
            return f"{response} {' '.join(map(str, args))}"
        return response
    return RESPONSES['en'].get(key, "Response not available")

# TTS thread safety
tts_lock = Lock()

def speak(text, language='en'):
    """Enhanced speak function with thread safety and better text cleaning"""
    if not TTS_AVAILABLE or not engine:
        print(f"JARVIS would say: {text}")
        return
    
    # Clean and validate text
    if not text or not isinstance(text, str):
        return
    
    # Advanced text cleaning to prevent garbled output
    cleaned_text = text.strip()
    if not cleaned_text:
        return
    
    # Remove problematic characters but preserve essential punctuation
    import re
    
    # Remove HTML tags if any
    cleaned_text = re.sub(r'<[^>]+>', '', cleaned_text)
    
    # Remove special characters that cause TTS issues but keep basic punctuation
    cleaned_text = re.sub(r'[^\w\s.,!?;:\-\'\"()]', ' ', cleaned_text)
    
    # Fix multiple spaces and normalize whitespace
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    
    # Ensure text isn't too long (TTS can fail with very long text)
    if len(cleaned_text) > 500:
        cleaned_text = cleaned_text[:500] + "..."
    
    # Use thread lock to prevent concurrent TTS calls
    with tts_lock:
        try:
            # Stop any ongoing speech first
            try:
                engine.stop()
            except:
                pass
            
            # Set speech properties for clarity
            engine.setProperty('rate', 160)  # Slightly faster but clear
            engine.setProperty('volume', 0.8)  # Good volume level
            
            # Try to set a good voice
            voices = engine.getProperty('voices')
            if voices and len(voices) > 0:
                # Find the best voice (prefer system default or first available)
                best_voice = None
                
                # Look for specific good voices
                for voice in voices:
                    voice_name = voice.name.lower() if hasattr(voice, 'name') else ''
                    
                    # Prefer these voices if available
                    if any(keyword in voice_name for keyword in ['alex', 'samantha', 'victoria', 'karen']):
                        best_voice = voice.id
                        break
                    elif 'english' in voice_name and not best_voice:
                        best_voice = voice.id
                
                # Set the voice
                if best_voice:
                    engine.setProperty('voice', best_voice)
                elif len(voices) > 0:
                    engine.setProperty('voice', voices[0].id)  # Use first available
            
            # Speak the text with error handling
            print(f"🔊 Speaking: {cleaned_text[:50]}...")
            engine.say(cleaned_text)
            engine.runAndWait()
            
        except Exception as e:
            print(f"TTS error: {e}")
            # Fallback: just print the text
            print(f"JARVIS (text-only): {cleaned_text}")
            
            # Try to reinitialize engine for next time
            try:
                if engine:
                    engine.stop()
                # Note: engine is a global variable, reinitializing here
                import pyttsx3
                new_engine = pyttsx3.init()
                if new_engine:
                    print("🔄 TTS engine reinitialized")
            except Exception as reinit_error:
                print(f"TTS reinit failed: {reinit_error}")
                pass

def get_wifi_details():
    try:
        ssid = os.popen("netsh wlan show interfaces").read()
        ip = socket.gethostbyname(socket.gethostname())
        return f"Connected Wi-Fi details: {ssid[:200]}... and IP Address: {ip}"
    except:
        return "Unable to fetch Wi-Fi details."

def get_battery_status():
    battery = psutil.sensors_battery()
    if battery:
        plugged = "Plugged In" if battery.power_plugged else "Not Plugged In"
        return f"Battery at {battery.percent}% ({plugged})"
    return "Battery status not available."

def search_files(keyword, folder=BASE_DIR):
    matches = []
    for root, _, files in os.walk(folder):
        for file in files:
            if keyword.lower() in file.lower():
                matches.append(os.path.join(root, file))
    return matches if matches else ["No files found"]

def set_reminder(msg, delay):
    def reminder_task():
        time.sleep(delay)
        speak(f"Reminder: {msg}")
    Thread(target=reminder_task).start()
    return f"Reminder set for {delay} seconds from now"

def control_volume(action):
    if not PYCAW_AVAILABLE:
        return "Volume control not available - pycaw not installed properly"
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
    volume = interface.QueryInterface(IAudioEndpointVolume)
    if action == "mute":
        volume.SetMute(1, None)
        return "Volume muted"
    elif action == "unmute":
        volume.SetMute(0, None)
        return "Volume unmuted"
    elif action == "increase":
        current = volume.GetMasterVolumeLevelScalar()
        volume.SetMasterVolumeLevelScalar(min(1.0, current + 0.1), None)
        return "Volume increased"
    elif action == "decrease":
        current = volume.GetMasterVolumeLevelScalar()
        volume.SetMasterVolumeLevelScalar(max(0.0, current - 0.1), None)
        return "Volume decreased"
    return "Unknown volume command"

# Create file
def create_file(path, filename):
    full_path = os.path.join(path, filename)
    with open(full_path, 'w') as f:
        pass
    return f"File {filename} created at {path}"

# Rename file
def rename_file(old, new):
    os.rename(old, new)
    return f"Renamed {old} to {new}"

# Delete file (to recycle bin)
def delete_file(path):
    send2trash.send2trash(path)
    return f"{path} moved to Recycle Bin"

# Find file
def find_file(directory, filename):
    for root, dirs, files in os.walk(directory):
        if filename in files:
            return os.path.join(root, filename)
    return None

# Open file
def open_file(path):
    os.startfile(path)
    return f"Opened {path}"
def take_screenshot():
    """Take a screenshot of the screen"""
    try:
        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save to Desktop if possible, otherwise to current directory
        try:
            desktop_path = Path.home() / 'Desktop'
            if desktop_path.exists():
                filename = desktop_path / f"JARVIS_screenshot_{timestamp}.png"
            else:
                filename = BASE_DIR / f"JARVIS_screenshot_{timestamp}.png"
        except:
            filename = BASE_DIR / f"JARVIS_screenshot_{timestamp}.png"
        
        # Take screenshot using pyautogui
        screenshot = pyautogui.screenshot()
        screenshot.save(str(filename))
        
        # Show notification if available
        if NOTIFICATIONS_AVAILABLE:
            show_notification("Screenshot Captured", f"Screenshot saved as {filename.name}")
        
        return f"Screenshot captured successfully and saved as {filename.name}, sir."
        
    except Exception as e:
        return f"Error taking screenshot: {str(e)}, sir."

def capture_photo():
    """Enhanced photo capture with cross-platform support"""
    if not CV2_AVAILABLE:
        return "Photo capture not available - OpenCV not installed. Please install opencv-python."
    
    try:
        # Try different camera indices for cross-platform compatibility
        camera_indices = [0, 1, 2]  # Try multiple camera sources
        cam = None
        
        for index in camera_indices:
            try:
                cam = cv2.VideoCapture(index)
                if cam.isOpened():
                    # Test if camera is working
                    ret, test_frame = cam.read()
                    if ret and test_frame is not None:
                        break
                    else:
                        cam.release()
                        cam = None
            except:
                if cam:
                    cam.release()
                cam = None
                continue
        
        if cam is None:
            return "No camera found or camera is not accessible, sir."
        
        # Set camera properties for better quality
        cam.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
        cam.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)
        cam.set(cv2.CAP_PROP_FPS, 30)
        
        # Capture frame
        ret, frame = cam.read()
        
        if ret and frame is not None:
            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Save to Desktop if possible, otherwise to current directory
            try:
                desktop_path = Path.home() / 'Desktop'
                if desktop_path.exists():
                    filename = desktop_path / f"JARVIS_photo_{timestamp}.jpg"
                else:
                    filename = BASE_DIR / f"JARVIS_photo_{timestamp}.jpg"
            except:
                filename = BASE_DIR / f"JARVIS_photo_{timestamp}.jpg"
            
            # Save the image
            success = cv2.imwrite(str(filename), frame)
            cam.release()
            
            if success:
                # Show notification if available
                if NOTIFICATIONS_AVAILABLE:
                    show_notification("Photo Captured", f"Photo saved as {filename.name}")
                return f"Photo captured successfully and saved as {filename.name}, sir."
            else:
                return "Failed to save the captured photo, sir."
        else:
            cam.release()
            return "Failed to capture photo from camera, sir."
            
    except Exception as e:
        if 'cam' in locals() and cam:
            cam.release()
        return f"Error capturing photo: {str(e)}, sir."

# Enhanced File Operations
def get_available_drives():
    """Get all available drives on Windows"""
    drives = []
    if IS_WINDOWS:
        import string
        for letter in string.ascii_uppercase:
            drive_path = Path(f"{letter}:/")
            if drive_path.exists():
                drives.append(drive_path)
    return drives

def find_files_and_folders_enhanced(name, search_path=None, max_results=20):
    """Enhanced file search across all drives and common locations"""
    matches = []
    searched_locations = []
    
    print(f"🔍 Searching for '{name}'...")
    
    # Define search locations based on OS
    if IS_WINDOWS:
        # Get all available drives
        drives = get_available_drives()
        search_locations = []
        
        # Priority locations first (user directories)
        priority_locations = [
            Path.home() / 'Desktop',
            Path.home() / 'Documents',
            Path.home() / 'Downloads',
            Path.home() / 'Pictures',
            Path.home() / 'Music',
            Path.home() / 'Videos',
            Path.home(),
        ]
        
        # Add drive root directories
        for drive in drives:
            search_locations.append(drive)
            # Add common folders in each drive
            common_folders = ['Users', 'Program Files', 'Program Files (x86)', 'Windows']
            for folder in common_folders:
                folder_path = drive / folder
                if folder_path.exists():
                    search_locations.append(folder_path)
        
        # Combine priority locations with drive locations
        search_locations = priority_locations + search_locations
        
    elif IS_MACOS:
        search_locations = [
            Path.home() / 'Desktop',
            Path.home() / 'Documents', 
            Path.home() / 'Downloads',
            Path.home() / 'Pictures',
            Path.home() / 'Music',
            Path.home() / 'Movies',
            Path.home(),
            Path('/Applications'),
            Path('/Users'),
            Path('/System/Applications'),
            Path('/'),
        ]
    else:  # Linux
        search_locations = [
            Path.home() / 'Desktop',
            Path.home() / 'Documents',
            Path.home() / 'Downloads',
            Path.home() / 'Pictures',
            Path.home() / 'Music',
            Path.home() / 'Videos',
            Path.home(),
            Path('/usr/bin'),
            Path('/usr/local/bin'),
            Path('/opt'),
            Path('/'),
        ]
    
    if search_path:
        search_locations.insert(0, Path(search_path))
    
    # Search with progress indication
    for i, search_dir in enumerate(search_locations):
        if len(matches) >= max_results:
            break
            
        try:
            if not search_dir.exists():
                continue
            
            searched_locations.append(str(search_dir))
            print(f"🔍 Searching in: {search_dir}")
            
            # Search in current directory first (non-recursive for speed)
            try:
                for item in search_dir.iterdir():
                    if len(matches) >= max_results:
                        break
                    if name.lower() in item.name.lower():
                        matches.append({
                            'path': str(item),
                            'name': item.name,
                            'type': 'folder' if item.is_dir() else 'file',
                            'size': item.stat().st_size if item.is_file() else 0,
                            'location': str(search_dir)
                        })
            except (PermissionError, OSError):
                continue
            
            # For user directories, search recursively but with depth limit
            if any(user_dir in str(search_dir).lower() for user_dir in ['desktop', 'documents', 'downloads', 'pictures', 'music', 'videos']):
                try:
                    for item in search_dir.rglob(f"*{name}*"):
                        if len(matches) >= max_results:
                            break
                        # Skip hidden files and system directories
                        if any(part.startswith('.') for part in item.parts):
                            continue
                        # Limit search depth to avoid performance issues
                        if len(item.parts) - len(search_dir.parts) > 3:
                            continue
                        matches.append({
                            'path': str(item),
                            'name': item.name,
                            'type': 'folder' if item.is_dir() else 'file',
                            'size': item.stat().st_size if item.is_file() else 0,
                            'location': str(search_dir)
                        })
                except (PermissionError, OSError):
                    continue
                    
        except (PermissionError, OSError):
            continue
    
    print(f"✅ Search complete. Found {len(matches)} matches.")
    return matches

def find_files_and_folders(name, search_path=None):
    """Legacy function - calls enhanced version for backward compatibility"""
    return find_files_and_folders_enhanced(name, search_path)

def handle_file_search_command(command_lower, original_command):
    """Handle file search commands with user interaction"""
    # Extract file name from command
    file_name = None
    search_terms = ['find file', 'search file', 'locate file', 'look for file', 'फाइल खोजो', 'ફાઇલ શોધો']
    
    for term in search_terms:
        if term in command_lower:
            file_name = command_lower.replace(term, '').strip()
            break
    
    if not file_name:
        return "Please specify what file you'd like me to search for, sir. For example: 'find file document.pdf'"
    
    # Search for files
    matches = find_files_and_folders_enhanced(file_name, max_results=10)
    
    if not matches:
        return f"Sorry sir, I couldn't find any files matching '{file_name}'. I searched across all available drives and common locations."
    
    # Format results with drive/location information
    result = f"I found {len(matches)} file(s) matching '{file_name}', sir:\n\n"
    
    for i, match in enumerate(matches, 1):
        file_type = "📁" if match['type'] == 'folder' else "📄"
        size_info = ""
        if match['type'] == 'file' and match['size'] > 0:
            size_mb = match['size'] / (1024 * 1024)
            if size_mb < 1:
                size_info = f" ({match['size']} bytes)"
            else:
                size_info = f" ({size_mb:.1f} MB)"
        
        # Show drive/location info
        location = match.get('location', '')
        if IS_WINDOWS:
            drive = match['path'][:3] if len(match['path']) > 2 else ''
            result += f"{i}. {file_type} {match['name']}{size_info}\n   📍 Location: {drive} - {match['path']}\n\n"
        else:
            result += f"{i}. {file_type} {match['name']}{size_info}\n   📍 Location: {match['path']}\n\n"
    
    # Store matches for potential opening
    global last_search_results
    last_search_results = matches
    
    result += "Would you like me to open any of these files, sir? Just say 'open file 1' or 'open file [number]' to open a specific file."
    
    return result

# Global variable to store last search results
last_search_results = []

# Song Recognition System
class SongRecognizer:
    def __init__(self):
        self.is_recording = False
        self.audio_data = []
        self.sample_rate = 44100
        self.chunk_size = 1024
        
        # Popular song database (simplified - in real implementation, you'd use a proper music database)
        self.song_database = {
            # Popular songs with their characteristic features (simplified)
            "shape of you": {
                "artist": "Ed Sheeran",
                "keywords": ["shape", "you", "love", "body", "crazy"],
                "youtube_search": "Ed Sheeran Shape of You official video"
            },
            "blinding lights": {
                "artist": "The Weeknd", 
                "keywords": ["blinding", "lights", "feel", "touch"],
                "youtube_search": "The Weeknd Blinding Lights official video"
            },
            "bad guy": {
                "artist": "Billie Eilish",
                "keywords": ["bad", "guy", "might", "seduce"],
                "youtube_search": "Billie Eilish bad guy official video"
            },
            "someone like you": {
                "artist": "Adele",
                "keywords": ["someone", "like", "you", "find", "love"],
                "youtube_search": "Adele Someone Like You official video"
            },
            "bohemian rhapsody": {
                "artist": "Queen",
                "keywords": ["bohemian", "rhapsody", "mama", "killed", "man"],
                "youtube_search": "Queen Bohemian Rhapsody official video"
            },
            "imagine": {
                "artist": "John Lennon",
                "keywords": ["imagine", "heaven", "hell", "peace"],
                "youtube_search": "John Lennon Imagine official video"
            },
            "hotel california": {
                "artist": "Eagles",
                "keywords": ["hotel", "california", "dark", "highway"],
                "youtube_search": "Eagles Hotel California official video"
            },
            "let it be": {
                "artist": "The Beatles",
                "keywords": ["let", "it", "be", "mother", "mary"],
                "youtube_search": "The Beatles Let It Be official video"
            },
            "rolling in the deep": {
                "artist": "Adele",
                "keywords": ["rolling", "deep", "fire", "heart"],
                "youtube_search": "Adele Rolling in the Deep official video"
            },
            "perfect": {
                "artist": "Ed Sheeran",
                "keywords": ["perfect", "tonight", "beautiful", "wonderful"],
                "youtube_search": "Ed Sheeran Perfect official video"
            }
        }
    
    def record_audio(self, duration=10):
        """Record audio from microphone"""
        if not AUDIO_PROCESSING_AVAILABLE:
            return None, "Audio processing libraries not available"
        
        try:
            print(f"🎤 Recording for {duration} seconds...")
            
            # Initialize PyAudio
            p = pyaudio.PyAudio()
            
            # Open stream
            stream = p.open(
                format=pyaudio.paInt16,
                channels=1,
                rate=self.sample_rate,
                input=True,
                frames_per_buffer=self.chunk_size
            )
            
            frames = []
            for _ in range(0, int(self.sample_rate / self.chunk_size * duration)):
                data = stream.read(self.chunk_size)
                frames.append(data)
            
            # Stop and close stream
            stream.stop_stream()
            stream.close()
            p.terminate()
            
            # Convert to numpy array
            audio_data = np.frombuffer(b''.join(frames), dtype=np.int16)
            audio_data = audio_data.astype(np.float32) / 32768.0  # Normalize
            
            print("✅ Recording completed!")
            return audio_data, "Recording successful"
            
        except Exception as e:
            return None, f"Recording error: {str(e)}"
    
    def extract_audio_features(self, audio_data):
        """Extract features from audio for song recognition"""
        try:
            # Extract basic audio features
            features = {}
            
            # Tempo detection
            tempo, _ = librosa.beat.beat_track(y=audio_data, sr=self.sample_rate)
            features['tempo'] = tempo
            
            # Pitch/frequency analysis
            pitches, magnitudes = librosa.piptrack(y=audio_data, sr=self.sample_rate)
            features['avg_pitch'] = np.mean(pitches[pitches > 0]) if np.any(pitches > 0) else 0
            
            # Spectral features
            spectral_centroids = librosa.feature.spectral_centroid(y=audio_data, sr=self.sample_rate)[0]
            features['spectral_centroid'] = np.mean(spectral_centroids)
            
            # Zero crossing rate
            zcr = librosa.feature.zero_crossing_rate(audio_data)[0]
            features['zcr'] = np.mean(zcr)
            
            # MFCC features (simplified)
            mfccs = librosa.feature.mfcc(y=audio_data, sr=self.sample_rate, n_mfcc=13)
            features['mfccs'] = np.mean(mfccs, axis=1)
            
            return features
            
        except Exception as e:
            print(f"Feature extraction error: {e}")
            return {}
    
    def recognize_song_from_lyrics(self, text_input):
        """Recognize song from sung/hummed lyrics or text"""
        text_lower = text_input.lower()
        
        # Score each song based on keyword matches
        song_scores = {}
        
        for song_title, song_info in self.song_database.items():
            score = 0
            
            # Check for exact title match
            if song_title in text_lower:
                score += 100
            
            # Check for artist name
            if song_info["artist"].lower() in text_lower:
                score += 50
            
            # Check for keyword matches
            for keyword in song_info["keywords"]:
                if keyword in text_lower:
                    score += 10
            
            if score > 0:
                song_scores[song_title] = score
        
        # Return best match
        if song_scores:
            best_song = max(song_scores, key=song_scores.get)
            return best_song, self.song_database[best_song]
        
        return None, None
    
    def search_and_play_song(self, song_title, song_info):
        """Search and play song on YouTube"""
        try:
            search_query = song_info["youtube_search"]
            youtube_url = f"https://www.youtube.com/results?search_query={search_query.replace(' ', '+')}"
            
            webbrowser.open(youtube_url)
            
            return f"🎵 Found it! Playing '{song_title}' by {song_info['artist']} on YouTube, sir."
            
        except Exception as e:
            return f"Error playing song: {str(e)}"

# Global song recognizer instance
song_recognizer = SongRecognizer() if AUDIO_PROCESSING_AVAILABLE else None

def handle_open_file_from_search(command_lower):
    """Handle opening files from previous search results"""
    global last_search_results
    
    if not last_search_results:
        return "No recent file search results found, sir. Please search for files first using 'find file [filename]'."
    
    # Extract file number
    import re
    number_match = re.search(r'open file (\d+)|open (\d+)', command_lower)
    
    if number_match:
        file_number = int(number_match.group(1) or number_match.group(2))
        
        if 1 <= file_number <= len(last_search_results):
            selected_file = last_search_results[file_number - 1]
            result = open_file_or_folder(selected_file['path'])
            
            # Clear search results after opening
            last_search_results = []
            
            return f"Opening {selected_file['name']}, sir. {result}"
        else:
            return f"Invalid file number, sir. Please choose a number between 1 and {len(last_search_results)}."
    
    return "Please specify which file number to open, sir. For example: 'open file 1' or 'open file 3'."

def handle_song_recognition(command_lower, original_command):
    """Handle song recognition from singing/humming"""
    global song_recognizer
    
    if not AUDIO_PROCESSING_AVAILABLE or not song_recognizer:
        return "Sorry sir, audio processing is not available. Please install the required audio libraries: pip install pyaudio librosa numpy scipy soundfile"
    
    # Check if user wants to sing/hum
    if any(phrase in command_lower for phrase in ['listen to me sing', 'i will sing', 'let me sing', 'recognize this song', 'what song is this']):
        return start_song_recording()
    
    # Check if user is providing lyrics directly
    elif any(phrase in command_lower for phrase in ['play song', 'find song', 'song with lyrics']):
        # Extract potential lyrics from command
        lyrics_indicators = ['play song', 'find song', 'song with lyrics', 'song that goes', 'lyrics']
        
        lyrics_text = original_command
        for indicator in lyrics_indicators:
            if indicator in command_lower:
                lyrics_text = original_command.split(indicator, 1)[-1].strip()
                break
        
        if lyrics_text and len(lyrics_text) > 3:
            return recognize_song_from_text(lyrics_text)
        else:
            return "Please provide some lyrics or say 'let me sing' so I can listen to you, sir."
    
    return None

def start_song_recording():
    """Start recording user singing/humming"""
    global song_recognizer
    
    try:
        response = "🎤 I'm ready to listen, sir! Please start singing or humming the song. I'll record for 10 seconds..."
        
        # Record audio
        audio_data, status = song_recognizer.record_audio(duration=10)
        
        if audio_data is None:
            return f"Recording failed: {status}"
        
        # For now, we'll ask user to provide lyrics since audio recognition is complex
        return "🎵 Recording complete! Since audio recognition is complex, could you please tell me some lyrics from the song you just sang? For example: 'The song goes: shape of you, I'm in love with your body'"
        
    except Exception as e:
        return f"Error during recording: {str(e)}"

def recognize_song_from_text(lyrics_text):
    """Recognize song from provided lyrics"""
    global song_recognizer
    
    try:
        song_title, song_info = song_recognizer.recognize_song_from_lyrics(lyrics_text)
        
        if song_title and song_info:
            result = song_recognizer.search_and_play_song(song_title, song_info)
            return result
        else:
            # If no match found, search YouTube with the lyrics
            search_query = lyrics_text.replace(' ', '+')
            youtube_url = f"https://www.youtube.com/results?search_query={search_query}+lyrics"
            webbrowser.open(youtube_url)
            return f"🎵 I couldn't identify the exact song, but I've searched YouTube for '{lyrics_text}' to help you find it, sir."
    
    except Exception as e:
        return f"Error recognizing song: {str(e)}"

def open_file_or_folder(path):
    """Open file or folder using system default application - cross-platform"""
    try:
        path_obj = Path(path)
        if path_obj.exists():
            if IS_WINDOWS:
                # Windows - use start command
                if path_obj.is_dir():
                    result = subprocess.run(['explorer', str(path_obj)], capture_output=True, text=True)
                else:
                    result = subprocess.run(['start', '', str(path_obj)], shell=True, capture_output=True, text=True)
                
                if result.returncode == 0:
                    return f"Successfully opened: {path_obj.name}"
                else:
                    return f"Error opening: {result.stderr}"
            
            elif IS_MACOS:
                # macOS - use open command
                result = subprocess.run(['open', str(path_obj)], capture_output=True, text=True)
                if result.returncode == 0:
                    return f"Successfully opened: {path_obj.name}"
                else:
                    return f"Error opening: {result.stderr}"
            
            else:  # Linux
                # Linux - use xdg-open
                try:
                    result = subprocess.run(['xdg-open', str(path_obj)], capture_output=True, text=True)
                    if result.returncode == 0:
                        return f"Successfully opened: {path_obj.name}"
                    else:
                        # Fallback to other methods
                        if path_obj.is_dir():
                            subprocess.run(['nautilus', str(path_obj)], capture_output=True)
                        else:
                            subprocess.run(['gedit', str(path_obj)], capture_output=True)
                        return f"Successfully opened: {path_obj.name}"
                except:
                    return f"Unable to open {path_obj.name} - no suitable application found"
        else:
            return f"Path not found: {path}"
    except Exception as e:
        return f"Error opening {path}: {str(e)}"

def delete_file_or_folder(path, permanent=False):
    """Delete file or folder (to trash by default) - macOS optimized"""
    try:
        path_obj = Path(path)
        if not path_obj.exists():
            return f"❌ Path not found: {path}, sir."
        
        item_name = path_obj.name
        item_type = "folder" if path_obj.is_dir() else "file"
        
        if permanent:
            if path_obj.is_dir():
                shutil.rmtree(str(path_obj))
                return f"✅ Permanently deleted {item_type} '{item_name}', sir."
            else:
                path_obj.unlink()
                return f"✅ Permanently deleted {item_type} '{item_name}', sir."
        else:
            # Try multiple methods to move to trash
            try:
                # Method 1: Use send2trash library
                if FILE_OPS_AVAILABLE:
                    import send2trash
                    send2trash.send2trash(str(path_obj))
                    return f"✅ Moved {item_type} '{item_name}' to trash, sir."
                else:
                    raise ImportError("send2trash not available")
            except Exception:
                try:
                    # Method 2: Use macOS osascript
                    script = f'tell application "Finder" to delete POSIX file "{str(path_obj)}"'
                    result = subprocess.run(['osascript', '-e', script], capture_output=True, text=True, timeout=10)
                    if result.returncode == 0:
                        return f"✅ Moved {item_type} '{item_name}' to trash, sir."
                    else:
                        # Method 3: Manual move to trash folder
                        trash_path = os.path.expanduser('~/.Trash')
                        if os.path.exists(trash_path):
                            import time
                            timestamp = int(time.time())
                            trash_item_path = os.path.join(trash_path, f"{item_name}_{timestamp}")
                            shutil.move(str(path_obj), trash_item_path)
                            return f"✅ Moved {item_type} '{item_name}' to trash, sir."
                        else:
                            return f"❌ Could not access trash. Use 'delete permanently' if needed, sir."
                except Exception as e2:
                    return f"❌ Error moving to trash: {str(e2)}. Try 'delete permanently' instead, sir."
    except PermissionError:
        return f"❌ Permission denied: Cannot delete {path}, sir."
    except Exception as e:
        return f"❌ Error deleting {path}: {str(e)}, sir."

def rename_file_or_folder(old_path, new_name):
    """Rename file or folder"""
    try:
        if not os.path.exists(old_path):
            return f"Path not found: {old_path}"
        
        old_path = Path(old_path)
        new_path = old_path.parent / new_name
        
        if new_path.exists():
            return f"Name already exists: {new_name}"
        
        old_path.rename(new_path)
        return f"Renamed {old_path.name} to {new_name}"
    except Exception as e:
        return f"Error renaming: {str(e)}"

def create_folder(path):
    """Create a new folder"""
    try:
        # Ensure the path is absolute
        if not os.path.isabs(path):
            desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
            path = os.path.join(desktop_path, path)
        
        os.makedirs(path, exist_ok=True)
        folder_name = os.path.basename(path)
        return f"✅ Successfully created folder '{folder_name}' at {path}, sir."
    except PermissionError:
        return f"❌ Permission denied: Cannot create folder at {path}, sir."
    except Exception as e:
        return f"❌ Error creating folder: {str(e)}, sir."

def handle_create_file_command(command_lower, original_command):
    """Handle file creation with support for all file types"""
    try:
        # Extract file name from command
        file_name = command_lower
        for phrase in ['create file', 'make file', 'new file', 'फाइल बनाओ', 'ફાઇલ બનાવો']:
            file_name = file_name.replace(phrase, '').strip()
        
        if not file_name:
            return "Please specify a file name, sir. For example: 'create file practical' or 'create file document.pdf'"
        
        # Parse file name and extension - if no extension provided, default to .txt
        if '.' not in file_name:
            # Default to .txt extension for simple file names
            file_name = file_name + '.txt'
        
        name_parts = file_name.rsplit('.', 1)
        base_name = name_parts[0]
        extension = name_parts[1].lower()
        
        # Create in Desktop by default
        desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        file_path = os.path.join(desktop_path, file_name)
        
        # Check if file already exists
        if os.path.exists(file_path):
            return f"❌ File '{file_name}' already exists on Desktop, sir. Please choose a different name."
        
        # Create file based on extension
        success = create_file_by_type(file_path, base_name, extension)
        if success:
            return f"✅ Successfully created {extension.upper()} file '{file_name}' on Desktop, sir."
        else:
            return f"❌ Failed to create {extension.upper()} file '{file_name}', sir."
            
    except Exception as e:
        return f"❌ Error creating file: {str(e)}, sir."

def create_file_by_type(file_path, base_name, extension):
    """Create different types of files with appropriate content"""
    try:
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        if extension in ['txt', 'text']:
            # Plain text file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {base_name}\n\nCreated by JARVIS on {current_time}\n\nThis is a text document.\n")
        
        elif extension in ['md', 'markdown']:
            # Markdown file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {base_name}\n\n**Created by JARVIS** on {current_time}\n\n## Content\n\nThis is a markdown document.\n\n- Item 1\n- Item 2\n- Item 3\n")
        
        elif extension in ['html', 'htm']:
            # HTML file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{base_name}</title>
</head>
<body>
    <h1>{base_name}</h1>
    <p>Created by JARVIS on {current_time}</p>
    <p>This is an HTML document.</p>
</body>
</html>""")
        
        elif extension in ['css']:
            # CSS file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"""/* {base_name} - Created by JARVIS on {current_time} */

body {{
    font-family: Arial, sans-serif;
    margin: 0;
    padding: 20px;
    background-color: #f5f5f5;
}}

.container {{
    max-width: 800px;
    margin: 0 auto;
    background: white;
    padding: 20px;
    border-radius: 8px;
    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
}}
""")
        
        elif extension in ['js', 'javascript']:
            # JavaScript file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"""// {base_name} - Created by JARVIS on {current_time}

console.log('Hello from {base_name}!');

// Your JavaScript code here
function main() {{
    console.log('JARVIS created this file on {current_time}');
}}

main();
""")
        
        elif extension in ['py', 'python']:
            # Python file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"""#!/usr/bin/env python3
# {base_name} - Created by JARVIS on {current_time}

def main():
    print("Hello from {base_name}!")
    print("Created by JARVIS on {current_time}")

if __name__ == "__main__":
    main()
""")
        
        elif extension in ['json']:
            # JSON file
            import json
            data = {
                "name": base_name,
                "created_by": "JARVIS",
                "created_on": current_time,
                "description": "JSON file created by JARVIS",
                "data": {}
            }
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
        
        elif extension in ['xml']:
            # XML file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"""<?xml version="1.0" encoding="UTF-8"?>
<document>
    <metadata>
        <name>{base_name}</name>
        <created_by>JARVIS</created_by>
        <created_on>{current_time}</created_on>
    </metadata>
    <content>
        <description>XML file created by JARVIS</description>
    </content>
</document>""")
        
        elif extension in ['csv']:
            # CSV file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"Name,Value,Created By,Created On\n")
                f.write(f"{base_name},Sample Data,JARVIS,{current_time}\n")
        
        elif extension in ['pdf']:
            # PDF file (requires reportlab)
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                c = canvas.Canvas(file_path, pagesize=letter)
                c.drawString(100, 750, f"{base_name}")
                c.drawString(100, 730, f"Created by JARVIS on {current_time}")
                c.drawString(100, 700, "This is a PDF document created by JARVIS.")
                c.save()
            except ImportError:
                # Fallback: create empty PDF-like file
                with open(file_path, 'wb') as f:
                    # Simple PDF header (not a real PDF, but has .pdf extension)
                    f.write(b'%PDF-1.4\n')
                    f.write(f'% {base_name} - Created by JARVIS on {current_time}\n'.encode())
        
        elif extension in ['docx', 'doc']:
            # Word document (requires python-docx)
            try:
                from docx import Document
                
                doc = Document()
                doc.add_heading(base_name, 0)
                doc.add_paragraph(f'Created by JARVIS on {current_time}')
                doc.add_paragraph('This is a Word document created by JARVIS.')
                doc.save(file_path)
            except ImportError:
                # Fallback: create RTF file with .docx extension
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"{base_name}\n\nCreated by JARVIS on {current_time}\n\nThis is a document file.")
        
        else:
            # Generic file - create as text
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {base_name}\n\nCreated by JARVIS on {current_time}\n\nFile type: {extension.upper()}\n")
        
        return True
        
    except Exception as e:
        print(f"Error creating file: {e}")
        return False

def handle_delete_command(command_lower, original_command):
    """Handle delete file/folder commands"""
    try:
        # Extract file/folder name from command
        item_name = command_lower
        for phrase in ['delete file', 'remove file', 'delete folder', 'remove folder', 'फाइल डिलीट करो', 'फोल्डर डिलीट करो', 'ફાઇલ ડિલીટ કરો', 'ફોલ્ડર ડિલીટ કરો']:
            item_name = item_name.replace(phrase, '').strip()
        
        # Check for permanent deletion
        permanent = 'permanently' in original_command.lower() or 'permanent' in original_command.lower()
        
        if item_name:
            # Search for the file/folder first
            matches = find_files_and_folders_enhanced(item_name, max_results=5)
            if matches:
                if len(matches) == 1:
                    return delete_file_or_folder(matches[0]['path'], permanent=permanent)
                else:
                    # Multiple matches - show options
                    result = f"Found {len(matches)} items named '{item_name}':\n"
                    for i, match in enumerate(matches[:5], 1):
                        item_type = "📁" if match.get('type') == 'folder' else "📄"
                        result += f"{i}. {item_type} {match['name']} ({match['path']})\n"
                    result += "\nPlease be more specific or use the full path, sir."
                    return result
            else:
                return f"❌ Could not find '{item_name}' to delete, sir."
        else:
            return "Please specify what to delete, sir. For example: 'delete file test.pdf' or 'delete folder MyFolder'"
            
    except Exception as e:
        return f"❌ Error processing delete command: {str(e)}, sir."

def handle_restore_file_command(command_lower, original_command):
    """Handle file restoration from trash"""
    try:
        # Extract file name from command
        file_name = command_lower
        for phrase in ['restore file', 'restore', 'undelete file', 'recover file']:
            file_name = file_name.replace(phrase, '').strip()
        
        if not file_name:
            return "Please specify the file name to restore, sir. For example: 'restore file test.txt'"
        
        # Check trash folder
        trash_path = os.path.expanduser('~/.Trash')
        if not os.path.exists(trash_path):
            return "❌ Trash folder not found, sir."
        
        # Find files in trash matching the name
        matches = []
        for item in os.listdir(trash_path):
            if file_name.lower() in item.lower():
                matches.append(item)
        
        if not matches:
            return f"❌ Could not find '{file_name}' in trash, sir."
        
        if len(matches) == 1:
            # Restore the file
            trash_item_path = os.path.join(trash_path, matches[0])
            desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
            
            # Remove timestamp suffix if present
            original_name = matches[0]
            if '_' in original_name and original_name.split('_')[-1].isdigit():
                original_name = '_'.join(original_name.split('_')[:-1])
            
            restore_path = os.path.join(desktop_path, original_name)
            
            # Check if file already exists at destination
            counter = 1
            while os.path.exists(restore_path):
                name_parts = original_name.rsplit('.', 1)
                if len(name_parts) == 2:
                    restore_path = os.path.join(desktop_path, f"{name_parts[0]}_restored_{counter}.{name_parts[1]}")
                else:
                    restore_path = os.path.join(desktop_path, f"{original_name}_restored_{counter}")
                counter += 1
            
            shutil.move(trash_item_path, restore_path)
            restored_name = os.path.basename(restore_path)
            return f"✅ Successfully restored '{restored_name}' to Desktop, sir."
        
        else:
            # Multiple matches
            result = f"Found {len(matches)} items in trash matching '{file_name}':\n"
            for i, match in enumerate(matches[:5], 1):
                result += f"{i}. {match}\n"
            result += "\nPlease be more specific, sir."
            return result
            
    except Exception as e:
        return f"❌ Error restoring file: {str(e)}, sir."

def handle_open_file_command(file_path):
    """Open file with default application"""
    try:
        if not os.path.exists(file_path):
            return f"❌ File not found: {file_path}, sir."
        
        file_name = os.path.basename(file_path)
        
        if IS_MACOS:
            subprocess.run(['open', file_path])
        elif IS_WINDOWS:
            os.startfile(file_path)
        else:  # Linux
            subprocess.run(['xdg-open', file_path])
        
        return f"✅ Opening '{file_name}' with default application, sir."
        
    except Exception as e:
        return f"❌ Error opening file: {str(e)}, sir."

def copy_file_or_folder(source, destination):
    """Copy file or folder"""
    try:
        if not os.path.exists(source):
            return f"Source not found: {source}"
        
        if os.path.isdir(source):
            shutil.copytree(source, destination)
            return f"Copied folder from {source} to {destination}"
        else:
            shutil.copy2(source, destination)
            return f"Copied file from {source} to {destination}"
    except Exception as e:
        return f"Error copying: {str(e)}"

def move_file_or_folder(source, destination):
    """Move file or folder"""
    try:
        if not os.path.exists(source):
            return f"Source not found: {source}"
        
        shutil.move(source, destination)
        return f"Moved from {source} to {destination}"
    except Exception as e:
        return f"Error moving: {str(e)}"

def get_file_info(path):
    """Get detailed information about a file or folder"""
    try:
        if not os.path.exists(path):
            return f"Path not found: {path}"
        
        path_obj = Path(path)
        stat = path_obj.stat()
        
        info = {
            'name': path_obj.name,
            'path': str(path_obj.absolute()),
            'type': 'folder' if path_obj.is_dir() else 'file',
            'size': stat.st_size,
            'created': datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
        }
        
        if path_obj.is_file():
            info['mime_type'] = mimetypes.guess_type(str(path_obj))[0]
        
        return f"File info: {json.dumps(info, indent=2)}"
    except Exception as e:
        return f"Error getting file info: {str(e)}"

def list_directory_contents(path, show_hidden=False):
    """List contents of a directory - macOS optimized"""
    try:
        path_obj = Path(path)
        if not path_obj.exists():
            return f"Path not found: {path}"
        
        if not path_obj.is_dir():
            return f"Not a directory: {path}"
        
        contents = []
        
        for item in path_obj.iterdir():
            if not show_hidden and item.name.startswith('.'):
                continue
            
            try:
                stat_info = item.stat()
                contents.append({
                    'name': item.name,
                    'type': 'folder' if item.is_dir() else 'file',
                    'size': stat_info.st_size if item.is_file() else 0
                })
            except (PermissionError, OSError):
                # Skip items we can't access
                continue
        
        contents.sort(key=lambda x: (x['type'] == 'file', x['name'].lower()))
        
        if not contents:
            return f"No accessible items found in {path_obj.name}"
        
        result = f"Contents of {path_obj.name}:\n"
        for item in contents[:20]:  # Limit to 20 items
            if item['type'] == 'file':
                size_mb = item['size'] / (1024 * 1024)
                if size_mb >= 1:
                    size_str = f" ({size_mb:.1f} MB)"
                else:
                    size_str = f" ({item['size']} bytes)"
            else:
                size_str = ""
            result += f"  📁 {item['name']}{size_str}\n" if item['type'] == 'folder' else f"  📄 {item['name']}{size_str}\n"
        
        if len(contents) > 20:
            result += f"  ... and {len(contents) - 20} more items"
        
        return result
    except Exception as e:
        return f"Error listing directory: {str(e)}"

def test_system_permissions():
    """Test if we have necessary permissions for file operations"""
    test_results = {
        'desktop_access': False,
        'documents_access': False,
        'downloads_access': False,
        'home_access': False
    }
    
    try:
        # Test Desktop access
        desktop = Path.home() / 'Desktop'
        if desktop.exists():
            list(desktop.iterdir())
            test_results['desktop_access'] = True
    except:
        pass
    
    try:
        # Test Documents access
        documents = Path.home() / 'Documents'
        if documents.exists():
            list(documents.iterdir())
            test_results['documents_access'] = True
    except:
        pass
    
    try:
        # Test Downloads access
        downloads = Path.home() / 'Downloads'
        if downloads.exists():
            list(downloads.iterdir())
            test_results['downloads_access'] = True
    except:
        pass
    
    try:
        # Test Home access
        home = Path.home()
        list(home.iterdir())
        test_results['home_access'] = True
    except:
        pass
    
    return test_results

def show_notification(title, message, timeout=5):
    """Show system notification"""
    if NOTIFICATIONS_AVAILABLE:
        try:
            notification.notify(
                title=title,
                message=message,
                timeout=timeout,
                app_name="JARVIS AI Assistant"
            )
        except Exception as e:
            print(f"Notification error: {e}")

def search_website(site, query):
    """Enhanced website search functionality with intelligent query processing"""
    if not query:
        return WEBSITE_MAP.get(site, f"https://www.{site}.com")
    
    # Clean and process the query
    query = query.strip()
    
    # Special handling for different sites
    if site == 'wikipedia':
        # For Wikipedia, clean up common phrases
        query = query.replace('prophet mohammad s.a.w.', 'Muhammad')
        query = query.replace('prophet muhammad', 'Muhammad')
        query = query.replace('of ', '')  # Remove "of" for cleaner searches
        encoded_query = urllib.parse.quote_plus(query.replace(' ', '_'))
        return f"https://en.wikipedia.org/wiki/{encoded_query}"
    
    # URL encode the query
    encoded_query = urllib.parse.quote_plus(query)
    
    if site in SEARCH_PATTERNS:
        return SEARCH_PATTERNS[site].format(encoded_query)
    else:
        # Fallback to Google search for the site
        return f"https://www.google.com/search?q=site:{site}.com {encoded_query}"

def get_weather_info(city=None):
    """Get weather information for a city"""
    try:
        if not city:
            # Try to get location from IP
            response = requests.get('http://ipapi.co/json/', timeout=5)
            if response.status_code == 200:
                data = response.json()
                city = data.get('city', 'Unknown')
        
        # You can integrate with a weather API here
        return f"Weather information for {city} - Please check your preferred weather app or website."
    except:
        return "Unable to fetch weather information at the moment."

def get_news_headlines():
    """Get latest news headlines"""
    try:
        # You can integrate with a news API here
        return "Opening news website for latest headlines..."
    except:
        return "Unable to fetch news at the moment."

def control_system_settings(action):
    """Control various system settings with improved cross-platform support"""
    try:
        if IS_WINDOWS:
            if action == 'shutdown':
                # Try different Windows shutdown methods
                try:
                    subprocess.run(['shutdown', '/s', '/t', '10'], check=True)
                    return "System will shutdown in 10 seconds, sir."
                except:
                    subprocess.run(['shutdown', '-s', '-t', '10'], check=True)
                    return "System will shutdown in 10 seconds, sir."
            elif action == 'restart':
                try:
                    subprocess.run(['shutdown', '/r', '/t', '10'], check=True)
                    return "System will restart in 10 seconds, sir."
                except:
                    subprocess.run(['shutdown', '-r', '-t', '10'], check=True)
                    return "System will restart in 10 seconds, sir."
            elif action == 'sleep':
                try:
                    subprocess.run(['rundll32.exe', 'powrprof.dll,SetSuspendState', '0,1,0'], check=True)
                    return "Putting system to sleep, sir."
                except:
                    subprocess.run(['powercfg', '/hibernate', 'off'], check=True)
                    return "Putting system to sleep, sir."
        
        elif IS_MACOS:
            if action == 'shutdown':
                try:
                    # Try without sudo first (user shutdown)
                    subprocess.run(['osascript', '-e', 'tell app "System Events" to shut down'], check=True)
                    return "System will shutdown shortly, sir."
                except:
                    try:
                        subprocess.run(['shutdown', '-h', '+1'], check=True)
                        return "System will shutdown in 1 minute, sir."
                    except:
                        return "Unable to shutdown - please check system permissions, sir."
            elif action == 'restart':
                try:
                    subprocess.run(['osascript', '-e', 'tell app "System Events" to restart'], check=True)
                    return "System will restart shortly, sir."
                except:
                    try:
                        subprocess.run(['shutdown', '-r', '+1'], check=True)
                        return "System will restart in 1 minute, sir."
                    except:
                        return "Unable to restart - please check system permissions, sir."
            elif action == 'sleep':
                try:
                    subprocess.run(['pmset', 'sleepnow'], check=True)
                    return "Putting system to sleep, sir."
                except:
                    try:
                        subprocess.run(['osascript', '-e', 'tell app "System Events" to sleep'], check=True)
                        return "Putting system to sleep, sir."
                    except:
                        return "Unable to sleep - please check system permissions, sir."
        
        else:  # Linux
            if action == 'shutdown':
                commands = [
                    ['systemctl', 'poweroff'],
                    ['shutdown', '-h', 'now'],
                    ['poweroff']
                ]
                for cmd in commands:
                    try:
                        subprocess.run(cmd, check=True)
                        return "System will shutdown now, sir."
                    except:
                        continue
                return "Unable to shutdown - please check system permissions, sir."
            elif action == 'restart':
                commands = [
                    ['systemctl', 'reboot'],
                    ['shutdown', '-r', 'now'],
                    ['reboot']
                ]
                for cmd in commands:
                    try:
                        subprocess.run(cmd, check=True)
                        return "System will restart now, sir."
                    except:
                        continue
                return "Unable to restart - please check system permissions, sir."
            elif action == 'sleep':
                commands = [
                    ['systemctl', 'suspend'],
                    ['pm-suspend'],
                    ['dbus-send', '--system', '--print-reply', '--dest=org.freedesktop.UPower', '/org/freedesktop/UPower', 'org.freedesktop.UPower.Suspend']
                ]
                for cmd in commands:
                    try:
                        subprocess.run(cmd, check=True)
                        return "Putting system to sleep, sir."
                    except:
                        continue
                return "Unable to sleep - please check system permissions, sir."
        
        return f"System control action '{action}' completed, sir."
    except Exception as e:
        return f"Error performing {action}: {str(e)}"

def get_system_performance():
    """Get detailed system performance metrics"""
    try:
        cpu_percent = psutil.cpu_percent(interval=1)
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/' if not IS_WINDOWS else 'C:')
        
        performance = {
            'cpu_usage': cpu_percent,
            'memory_usage': memory.percent,
            'memory_available': round(memory.available / (1024**3), 2),
            'disk_free': round(disk.free / (1024**3), 2),
            'disk_usage': round((disk.used / disk.total) * 100, 2)
        }
        
        return f"System Performance: CPU {cpu_percent}%, Memory {memory.percent}%, Disk {performance['disk_usage']}% used"
    except Exception as e:
        return f"Error getting system performance: {str(e)}"

def manage_clipboard(action, text=None):
    """Manage clipboard operations"""
    try:
        if action == 'copy' and text:
            pyperclip.copy(text)
            return f"Copied to clipboard: {text[:50]}..."
        elif action == 'paste':
            content = pyperclip.paste()
            return f"Clipboard content: {content[:100]}..."
        elif action == 'clear':
            pyperclip.copy('')
            return "Clipboard cleared, sir."
        return "Invalid clipboard operation."
    except Exception as e:
        return f"Clipboard error: {str(e)}"

def set_timer(duration_minutes):
    """Set a timer for specified minutes"""
    try:
        duration_seconds = duration_minutes * 60
        
        def timer_task():
            time.sleep(duration_seconds)
            if NOTIFICATIONS_AVAILABLE:
                show_notification("JARVIS Timer", f"Timer for {duration_minutes} minutes has finished!")
            speak(f"Timer for {duration_minutes} minutes has finished, sir.")
        
        Thread(target=timer_task).start()
        return f"Timer set for {duration_minutes} minutes, sir."
    except Exception as e:
        return f"Error setting timer: {str(e)}"

def calculate_expression(expression):
    """Safely calculate mathematical expressions"""
    try:
        # Remove any potentially dangerous characters
        safe_chars = set('0123456789+-*/.() ')
        if not all(c in safe_chars for c in expression):
            return "Invalid characters in expression."
        
        # Use eval safely for basic math
        result = eval(expression)
        return f"The result is: {result}"
    except Exception as e:
        return f"Error calculating: {str(e)}"

def get_random_fact():
    """Get a random interesting fact"""
    facts = [
        "The human brain contains approximately 86 billion neurons.",
        "A group of flamingos is called a 'flamboyance'.",
        "Honey never spoils - archaeologists have found edible honey in ancient Egyptian tombs.",
        "The shortest war in history lasted only 38-45 minutes between Britain and Zanzibar in 1896.",
        "Octopuses have three hearts and blue blood.",
        "The Great Wall of China isn't visible from space with the naked eye.",
        "A single cloud can weigh more than a million pounds.",
        "Bananas are berries, but strawberries aren't.",
        "The human body contains about 37.2 trillion cells.",
        "Lightning strikes the Earth about 100 times per second."
    ]
    return random.choice(facts)

def get_wikipedia_summary(topic, sentences=2):
    """Get Wikipedia summary for a topic"""
    try:
        # Clean up the topic
        topic = topic.replace('prophet mohammad s.a.w.', 'Muhammad')
        topic = topic.replace('prophet muhammad', 'Muhammad')
        
        summary = wikipedia.summary(topic, sentences=sentences)
        return summary
    except wikipedia.exceptions.DisambiguationError as e:
        # If there are multiple options, use the first one
        try:
            summary = wikipedia.summary(e.options[0], sentences=sentences)
            return summary
        except:
            return f"Multiple topics found for '{topic}'. Please be more specific."
    except wikipedia.exceptions.PageError:
        return f"No Wikipedia page found for '{topic}'."
    except Exception as e:
        return f"Error retrieving Wikipedia information: {str(e)}"

def open_application_cross_platform(app_name):
    """Open application across different operating systems with fallbacks"""
    try:
        if app_name in APP_MAP:
            app_paths = APP_MAP[app_name] if isinstance(APP_MAP[app_name], list) else [APP_MAP[app_name]]
            
            for app_path in app_paths:
                try:
                    if IS_WINDOWS:
                        # Expand environment variables
                        app_path = os.path.expandvars(app_path)
                        
                        if app_path.endswith('.exe'):
                            if os.path.exists(app_path):
                                subprocess.Popen([app_path])
                                return f"Opening {app_name} application, sir."
                            else:
                                # Try to find in PATH
                                try:
                                    subprocess.Popen([os.path.basename(app_path)])
                                    return f"Opening {app_name} application, sir."
                                except:
                                    continue
                        else:
                            subprocess.Popen(['start', '', app_path], shell=True)
                            return f"Opening {app_name} application, sir."
                    
                    elif IS_MACOS:
                        if os.path.exists(app_path):
                            subprocess.run(['open', '-a', app_path])
                            return f"Opening {app_name} application, sir."
                        else:
                            continue
                    
                    else:  # Linux
                        try:
                            subprocess.Popen([app_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                            return f"Opening {app_name} application, sir."
                        except FileNotFoundError:
                            continue
                
                except Exception as e:
                    print(f"Failed to open {app_path}: {e}")
                    continue
            
            # If all paths failed, try generic approach
            if IS_WINDOWS:
                try:
                    subprocess.Popen(['start', '', app_name], shell=True)
                    return f"Attempting to open {app_name}, sir."
                except:
                    pass
            elif IS_MACOS:
                try:
                    subprocess.run(['open', '-a', app_name])
                    return f"Attempting to open {app_name}, sir."
                except:
                    pass
            else:  # Linux
                try:
                    subprocess.Popen([app_name], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    return f"Attempting to open {app_name}, sir."
                except:
                    pass
            
            return f"Application '{app_name}' not found or cannot be opened, sir."
        else:
            return f"Application '{app_name}' not found in my database, sir."
    except Exception as e:
        return f"Error opening {app_name}: {str(e)}"

def get_system_info():
    """Get comprehensive system information"""
    try:
        info = {
            'os': platform.system(),
            'os_version': platform.version(),
            'architecture': platform.architecture()[0],
            'processor': platform.processor(),
            'hostname': platform.node(),
            'python_version': platform.python_version(),
            'cpu_count': psutil.cpu_count(),
            'memory_total': round(psutil.virtual_memory().total / (1024**3), 2),
            'memory_available': round(psutil.virtual_memory().available / (1024**3), 2),
            'disk_usage': round(psutil.disk_usage('/').free / (1024**3), 2) if not IS_WINDOWS else round(psutil.disk_usage('C:').free / (1024**3), 2)
        }
        return info
    except Exception as e:
        return f"Error getting system info: {str(e)}"

def normalize_multilingual_command(command):
    """Normalize multilingual commands to English equivalents"""
    # Hindi to English mappings
    hindi_mappings = {
        'समय क्या है': 'what time is it',
        'समय': 'time',
        'आज की तारीख': 'what is today date',
        'तारीख': 'date',
        'बैटरी स्टेटस': 'battery status',
        'बैटरी': 'battery',
        'मदद करो': 'help me',
        'मदद': 'help',
        'मजाक सुनाओ': 'tell me a joke',
        'मजाक': 'joke',
        'खोलो': 'open',
        'बंद करो': 'close',
        'चालू करो': 'start',
        'फाइल बनाओ': 'create file',
        'फोल्डर बनाओ': 'create folder',
        'फाइल डिलीट करो': 'delete file',
        'फोल्डर डिलीट करो': 'delete folder',
        'वाईफाई': 'wifi',
        'इंटरनेट': 'internet',
        'फोटो खींचो': 'take photo',
        'आवाज बढ़ाओ': 'volume up',
        'आवाज कम करो': 'volume down',
    }
    
    # Gujarati to English mappings
    gujarati_mappings = {
        'સમય શું છે': 'what time is it',
        'સમય': 'time',
        'આજની તારીખ': 'what is today date',
        'તારીખ': 'date',
        'બેટરી સ્ટેટસ': 'battery status',
        'બેટરી': 'battery',
        'મદદ કરો': 'help me',
        'મદદ': 'help',
        'મજાક કહો': 'tell me a joke',
        'મજાક': 'joke',
        'ખોલો': 'open',
        'બંધ કરો': 'close',
        'ચાલુ કરો': 'start',
        'ફાઇલ બનાવો': 'create file',
        'ફોલ્ડર બનાવો': 'create folder',
        'ફાઇલ ડિલીટ કરો': 'delete file',
        'ફોલ્ડર ડિલીટ કરો': 'delete folder',
        'વાઈફાઈ': 'wifi',
        'ઈન્ટરનેટ': 'internet',
        'ફોટો લો': 'take photo',
        'અવાજ વધારો': 'volume up',
        'અવાજ ઘટાડો': 'volume down',
    }
    
    # Apply mappings
    normalized = command.lower()
    
    # Apply Hindi mappings
    for hindi, english in hindi_mappings.items():
        if hindi in normalized:
            normalized = normalized.replace(hindi, english)
    
    # Apply Gujarati mappings
    for gujarati, english in gujarati_mappings.items():
        if gujarati in normalized:
            normalized = normalized.replace(gujarati, english)
    
    return normalized

def process_command(command, language='en'):
    """AI-First Command Processing - Natural Conversation Like Siri/Google Assistant"""
    if not command or not command.strip():
        return "I didn't hear anything, sir. Could you please repeat that?"
    
    original_command = command.strip()
    command_lower = command.lower().strip()
    
    print(f"🎤 Processing: '{original_command}'")
    
    # Detect language
    detected_language = detect_language(command)
    print(f"🌐 Language: {detected_language}")
    
    # Remove wake words if present
    wake_words = ['hey jarvis', 'jarvis', 'हे जार्विस', 'જાર્વિસ']
    for wake_word in wake_words:
        if command_lower.startswith(wake_word):
            command_lower = command_lower.replace(wake_word, '').strip()
            original_command = original_command[len(wake_word):].strip()
            break
    
    # If empty after wake word removal, greet
    if not command_lower:
        return "Hello! I'm JARVIS, your AI assistant. How may I help you today, sir?"
    
    # IMPROVED FILE OPERATIONS - HIGHEST PRIORITY
    if IMPROVED_FILE_OPS_AVAILABLE:
        file_result = command_processor.process_file_command(command_lower, original_command)
        if file_result:
            print(f"📁 File operation result: {file_result[:100]}...")
            return file_result
    
    # SONG RECOGNITION COMMANDS - HIGHEST PRIORITY (before AI processing)
    song_recognition_phrases = [
        'recognize this song', 'what song is this', 'identify this song', 
        'name this song', 'tell me this song', 'what is this song',
        'listen to this song', 'identify the song', 'recognize the song'
    ]
    
    if any(phrase in command_lower for phrase in song_recognition_phrases):
        print(f"🎵 Song recognition command detected: '{original_command}'")
        return "🎤 I'm ready to listen to the song playing in the background, sir! Please make sure the song is playing clearly and I'll try to identify it. You can also tell me some lyrics like 'the song goes: shape of you' to help me identify it better."
    
    # Singing/humming recognition commands
    singing_phrases = [
        'listen to me sing', 'i will sing', 'let me sing', 'i want to sing'
    ]
    
    if any(phrase in command_lower for phrase in singing_phrases):
        print(f"🎵 Singing recognition command detected: '{original_command}'")
        song_result = handle_song_recognition(command_lower, original_command)
        if song_result:
            return song_result
    
    # Lyrics-based song search
    lyrics_phrases = [
        'play song', 'find song', 'song with lyrics', 'song that goes', 'the song goes'
    ]
    
    if any(phrase in command_lower for phrase in lyrics_phrases):
        print(f"🎵 Lyrics-based song search detected: '{original_command}'")
        song_result = handle_song_recognition(command_lower, original_command)
        if song_result:
            return song_result
    
    # AI-FIRST APPROACH - Let AI handle most queries naturally
    if AI_AVAILABLE:
        try:
            # Check if it's a direct system command that needs immediate handling
            direct_system_commands = [
                'screenshot', 'screen shot', 'capture screen',
                'take photo', 'capture photo', 'take picture',
                'volume mute', 'mute volume', 'volume unmute', 'unmute volume',
                'volume up', 'volume down', 'increase volume', 'decrease volume',
                'create file', 'make file', 'new file', 'create folder', 'make folder', 'new folder',
                'delete file', 'remove file', 'delete folder', 'remove folder'
            ]
            
            is_direct_system_command = any(cmd in command_lower for cmd in direct_system_commands)
            
            if not is_direct_system_command:
                print(f"🤖 Sending to AI: '{original_command}'")
                ai_response = get_ai_response(original_command)
                if ai_response and len(ai_response.strip()) > 5:
                    print(f"✅ AI Response: {ai_response[:100]}...")
                    return ai_response
        except Exception as e:
            print(f"❌ AI processing error: {e}")
    
    # DIRECT SYSTEM COMMANDS - Only essential system operations
    
    # Screenshot
    if any(phrase in command_lower for phrase in ['screenshot', 'screen shot', 'capture screen']):
        result = take_screenshot()
        return result
    
    # Photo capture - Enhanced version that auto-opens photos
    if any(phrase in command_lower for phrase in ['take photo', 'capture photo', 'take picture', 'photo']):
        if IMPROVED_FILE_OPS_AVAILABLE:
            result = file_ops.take_photo_and_open()
            return result['message']
        else:
            result = capture_photo()
            return result
    
    # Volume control
    if 'volume' in command_lower:
        if any(word in command_lower for word in ['mute', 'silent']):
            control_volume('mute')
            return "Volume muted, sir."
        elif any(word in command_lower for word in ['unmute', 'turn on']):
            control_volume('unmute')
            return "Volume unmuted, sir."
        elif any(word in command_lower for word in ['up', 'increase', 'higher']):
            control_volume('increase')
            return "Volume increased, sir."
        elif any(word in command_lower for word in ['down', 'decrease', 'lower']):
            control_volume('decrease')
            return "Volume decreased, sir."
    
    # Website opening
    for site, url in WEBSITE_MAP.items():
        if f"open {site}" in command_lower or site in command_lower:
            webbrowser.open(url)
            return f"Opening {site.title()}, sir."
    
    # Search functionality
    if any(phrase in command_lower for phrase in ['search for', 'search', 'look up']):
        search_query = command_lower
        for phrase in ['search for', 'search', 'look up']:
            search_query = search_query.replace(phrase, '').strip()
        
        if search_query:
            if 'youtube' in command_lower:
                url = SEARCH_PATTERNS['youtube'].format(search_query.replace(' ', '+'))
                webbrowser.open(url)
                return f"Searching YouTube for '{search_query}', sir."
            else:
                url = SEARCH_PATTERNS['google'].format(search_query.replace(' ', '+'))
                webbrowser.open(url)
                return f"Searching Google for '{search_query}', sir."
    
    # Application opening
    if 'open' in command_lower:
        for app_name, app_paths in APP_MAP.items():
            if app_name in command_lower:
                success = False
                for app_path in app_paths:
                    try:
                        if IS_WINDOWS:
                            if app_path.endswith('.exe'):
                                subprocess.Popen(app_path, shell=True)
                            else:
                                subprocess.Popen(['start', app_path], shell=True)
                        elif IS_MACOS:
                            subprocess.Popen(['open', '-a', app_path])
                        else:  # Linux
                            subprocess.Popen([app_path])
                        success = True
                        break
                    except Exception:
                        continue
                
                if success:
                    return f"Opening {app_name.title()}, sir."
                else:
                    return f"I couldn't find {app_name} on your system, sir."
    
    # FALLBACK TO AI - Let AI handle everything else naturally
    # ENHANCED FILE SEARCH COMMANDS
    if any(term in command_lower for term in ['find file', 'search file', 'locate file', 'look for file', 'फाइल खोजो', 'ફાઇલ શોધો']):
        return handle_file_search_command(command_lower, original_command)
    
    # OPEN FILE FROM SEARCH RESULTS
    if any(term in command_lower for term in ['open file', 'open number']) and any(char.isdigit() for char in command_lower):
        return handle_open_file_from_search(command_lower)
    
    # FILE OPERATIONS COMMANDS
    if FILE_OPS_AVAILABLE or True:  # Enable file operations even without libraries (using built-in methods)
        
        # Create folder
        if any(phrase in command_lower for phrase in ['create folder', 'make folder', 'new folder', 'फोल्डर बनाओ', 'ફોલ્ડર બનાવો']):
            # Extract folder name from command
            folder_name = command_lower
            for phrase in ['create folder', 'make folder', 'new folder', 'फोल्डर बनाओ', 'ફોલ્ડર બનાવો']:
                folder_name = folder_name.replace(phrase, '').strip()
            
            if folder_name:
                # Create in Desktop by default
                desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
                folder_path = os.path.join(desktop_path, folder_name)
                return create_folder(folder_path)
            else:
                return "Please specify a folder name, sir. For example: 'create folder MyNewFolder'"
        
        # Create file
        if any(phrase in command_lower for phrase in ['create file', 'make file', 'new file', 'फाइल बनाओ', 'ફાઇલ બનાવો']):
            return handle_create_file_command(command_lower, original_command)

        # Create file (implicit) - handle commands like "create note.txt" or "create notes"
        # Only trigger if user didn't explicitly say folder/file keywords
        import re as _re
        if (
            _re.match(r'^(create|make|new)\s+\S+', command_lower)
            and not any(phrase in command_lower for phrase in ['create folder', 'make folder', 'new folder'])
            and not any(phrase in command_lower for phrase in ['create file', 'make file', 'new file'])
        ):
            # Extract the name after the first verb
            parts = command_lower.split(maxsplit=1)
            implicit_name = parts[1] if len(parts) > 1 else ''
            if implicit_name:
                synthetic_command = f"create file {implicit_name}"
                return handle_create_file_command(synthetic_command, original_command)
        
        # Delete file/folder
        if any(phrase in command_lower for phrase in ['delete file', 'remove file', 'delete folder', 'remove folder', 'फाइल डिलीट करो', 'फोल्डर डिलीट करो', 'ફાઇલ ડિલીટ કરો', 'ફોલ્ડર ડિલીટ કરો']):
            return handle_delete_command(command_lower, original_command)
        
        # Restore file from trash
        if any(phrase in command_lower for phrase in ['restore file', 'restore', 'undelete file', 'recover file']):
            return handle_restore_file_command(command_lower, original_command)
        
        # Open file after finding
        if any(phrase in command_lower for phrase in ['open file', 'launch file', 'run file']) and not any(char.isdigit() for char in command_lower):
            # Extract file name from command
            file_name = command_lower
            for phrase in ['open file', 'launch file', 'run file']:
                file_name = file_name.replace(phrase, '').strip()
            
            if file_name:
                # Search for the file first
                matches = find_files_and_folders_enhanced(file_name, max_results=5)
                if matches:
                    if len(matches) == 1:
                        return handle_open_file_command(matches[0]['path'])
                    else:
                        # Multiple matches - show options and store for numbered selection
                        global last_search_results
                        last_search_results = matches
                        result = f"Found {len(matches)} files named '{file_name}':\n"
                        for i, match in enumerate(matches[:10], 1):
                            result += f"{i}. {match['name']} ({match['path']})\n"
                        result += f"\nSay 'open number X' to open a specific file, sir."
                        return result
                else:
                    return f"❌ Could not find file '{file_name}', sir."
            else:
                return "Please specify the file name to open, sir. For example: 'open file test.pdf'"
        
        # Rename file/folder
        if any(phrase in command_lower for phrase in ['rename file', 'rename folder', 'rename']):
            # This is more complex - would need "rename X to Y" format
            if ' to ' in command_lower:
                parts = command_lower.split(' to ')
                if len(parts) == 2:
                    old_name = parts[0]
                    for phrase in ['rename file', 'rename folder', 'rename']:
                        old_name = old_name.replace(phrase, '').strip()
                    new_name = parts[1].strip()
                    
                    if old_name and new_name:
                        # Search for the file/folder first
                        matches = find_files_and_folders_enhanced(old_name, max_results=5)
                        if matches:
                            if len(matches) == 1:
                                return rename_file_or_folder(matches[0]['path'], new_name)
                            else:
                                result = f"Found {len(matches)} items named '{old_name}':\n"
                                for i, match in enumerate(matches[:5], 1):
                                    result += f"{i}. {match['name']} ({match['path']})\n"
                                result += "\nPlease be more specific, sir."
                                return result
                        else:
                            return f"Could not find '{old_name}' to rename, sir."
            
            return "Please use format: 'rename [filename] to [newname]', sir. For example: 'rename MyFile.txt to NewFile.txt'"
    
    # Song recognition commands are now handled at the top of the function
    
    if AI_AVAILABLE:
        try:
            print(f"🤖 AI Fallback for: '{original_command}'")
            ai_response = get_ai_response(original_command)
            if ai_response:
                return ai_response
        except Exception as e:
            print(f"❌ AI fallback error: {e}")
    
    # FINAL FALLBACK
    return f"I understand you said '{original_command}', but I'm not sure how to help with that, sir. Could you try rephrasing it or ask me something else?"


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/command', methods=['POST'])
def handle_command():
    command = request.json.get('command', '')
    
    # Detect language of the command
    detected_language = detect_language(command)
    
    # Process command with language context
    response = process_command(command, detected_language)
    
    # Start speaking the response in the detected language
    Thread(target=speak, args=(response, detected_language)).start()
    
    return jsonify({
        'response': response,
        'language': detected_language,
        'original_command': command
    })

@app.route('/api/speak', methods=['POST'])
def handle_speak():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'status': 'error', 'message': 'No data provided'})
        
        text = data.get('text', '').strip()
        language = data.get('language', 'en')
        
        if not text:
            return jsonify({'status': 'error', 'message': 'No text provided'})
        
        # Enhanced text validation and cleaning
        if not isinstance(text, str):
            return jsonify({'status': 'error', 'message': 'Invalid text format'})
        
        # Clean the text before processing
        import re
        cleaned_text = text.strip()
        
        # Remove HTML tags
        cleaned_text = re.sub(r'<[^>]+>', '', cleaned_text)
        
        # Remove problematic characters but keep essential punctuation
        cleaned_text = re.sub(r'[^\w\s.,!?;:\-\'\"()]', ' ', cleaned_text)
        
        # Fix multiple spaces
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        
        if not cleaned_text:
            return jsonify({'status': 'error', 'message': 'No valid text after cleaning'})
        
        # Limit text length to prevent TTS overload
        if len(cleaned_text) > 500:
            cleaned_text = cleaned_text[:500] + "..."
        
        print(f"🔊 TTS Request: {cleaned_text[:100]}...")
        
        # Check if TTS is available
        if not TTS_AVAILABLE or not engine:
            return jsonify({
                'status': 'error', 
                'message': 'TTS not available - using fallback',
                'fallback': True
            })
        
        # Start TTS in a separate thread with better error handling
        def safe_speak():
            try:
                speak(cleaned_text, language)
                print("✅ TTS completed successfully")
            except Exception as e:
                print(f"❌ TTS thread error: {e}")
        
        Thread(target=safe_speak, daemon=True).start()
        
        return jsonify({
            'status': 'speaking', 
            'language': language, 
            'text_length': len(cleaned_text),
            'original_length': len(text),
            'cleaned': True
        })
        
    except Exception as e:
        print(f"❌ TTS endpoint error: {e}")
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/api/translate', methods=['POST'])
def handle_translate():
    text = request.json.get('text', '')
    target_lang = request.json.get('target_lang', 'en')
    
    if TRANSLATION_AVAILABLE:
        translated = translate_text(text, target_lang)
        return jsonify({
            'original': text,
            'translated': translated,
            'target_language': target_lang
        })
    else:
        return jsonify({
            'error': 'Translation not available',
            'original': text
        })

@app.route('/api/system-status', methods=['GET'])
def handle_system_status():
    permissions = test_system_permissions()
    return jsonify({
        'permissions': permissions,
        'file_ops_available': FILE_OPS_AVAILABLE,
        'translation_available': TRANSLATION_AVAILABLE,
        'tts_available': TTS_AVAILABLE,
        'cv2_available': CV2_AVAILABLE
    })

@app.route('/api/capture-photo', methods=['POST'])
def handle_photo_capture():
    """Handle photo capture requests"""
    try:
        result = capture_photo()
        return jsonify({
            'status': 'success',
            'message': result
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f"Photo capture failed: {str(e)}"
        })

@app.route('/api/take-screenshot', methods=['POST'])
def handle_screenshot():
    """Handle screenshot requests"""
    try:
        result = take_screenshot()
        return jsonify({
            'status': 'success',
            'message': result
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f"Screenshot failed: {str(e)}"
        })

@app.route('/api/ai-query', methods=['POST'])
def api_ai_query():
    """API endpoint for AI-powered queries"""
    try:
        query = request.json.get('query', '')
        if not query:
            return jsonify({
                'status': 'error',
                'message': 'No query provided'
            })
        
        if AI_AVAILABLE:
            response = get_ai_response(query)
            return jsonify({
                'status': 'success',
                'message': response,
                'ai_powered': True
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'AI assistant not available',
                'ai_powered': False
            })
            
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'AI query failed: {str(e)}'
        })

@app.route('/api/song-recognition', methods=['POST'])
def api_song_recognition():
    """API endpoint for song recognition"""
    try:
        data = request.json
        action = data.get('action', '')
        lyrics = data.get('lyrics', '')
        
        if not AUDIO_PROCESSING_AVAILABLE:
            return jsonify({
                'status': 'error',
                'message': 'Audio processing not available. Please install required libraries.'
            })
        
        if action == 'recognize_lyrics' and lyrics:
            song_title, song_info = song_recognizer.recognize_song_from_lyrics(lyrics)
            
            if song_title and song_info:
                return jsonify({
                    'status': 'success',
                    'song_title': song_title,
                    'artist': song_info['artist'],
                    'youtube_search': song_info['youtube_search'],
                    'message': f"Found: {song_title} by {song_info['artist']}"
                })
            else:
                return jsonify({
                    'status': 'not_found',
                    'message': 'Song not found in database',
                    'youtube_search': f"https://www.youtube.com/results?search_query={lyrics.replace(' ', '+')}"
                })
        
        elif action == 'start_recording':
            return jsonify({
                'status': 'success',
                'message': 'Recording started. Please sing or hum the song.',
                'duration': 10
            })
        
        else:
            return jsonify({
                'status': 'error',
                'message': 'Invalid action or missing lyrics'
            })
            
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        })

@app.route('/api/test-file-operation', methods=['POST'])
def test_file_operation():
    """Test file operations with a simple file creation/deletion"""
    try:
        test_file = Path.home() / 'Desktop' / 'jarvis_test.txt'
        
        # Create test file
        test_file.write_text("JARVIS file operation test")
        
        # Check if file exists
        if test_file.exists():
            # Delete test file
            test_file.unlink()
            return jsonify({
                'success': True,
                'message': 'File operations working correctly, sir!'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'File creation failed - permission issues detected'
            })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'File operation test failed: {str(e)}'
        })

if __name__ == '__main__':
    # Bind to localhost on port 8888 explicitly
    app.run(host='127.0.0.1', port=8888, debug=True)
